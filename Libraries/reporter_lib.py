# reporter_lib.py

import sys
import os
import datetime
import io
import re
import logging
import glob
from docx import Document
from docx.shared import Inches, Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION, WD_ORIENT
import matplotlib.pyplot as plt
import seaborn as sns
import numpy as np
import pandas as pd
import openpyxl
from openpyxl.worksheet.filters import FilterColumn, Filters
from typing import List, Dict, Tuple
from sklearn.metrics import confusion_matrix, accuracy_score, precision_recall_fscore_support
from sklearn.metrics.pairwise import cosine_similarity

# Additional imports for SVG conversion
try:
    from cairosvg import svg2png
    HAS_CAIROSVG = True
except ImportError:
    HAS_CAIROSVG = False
    
try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

# Function to find the next available report directory
def get_next_report_dir():
    """
    Find the next available report directory in the format Report_XX
    """
    base_dir = "Results"
    os.makedirs(base_dir, exist_ok=True)
    
    # Find existing report directories
    existing_dirs = []
    for item in os.listdir(base_dir):
        if os.path.isdir(os.path.join(base_dir, item)) and item.startswith("Report_"):
            try:
                num = int(item.split("_")[1])
                existing_dirs.append(num)
            except ValueError:
                continue
    
    # Determine the next number
    next_num = 1
    if existing_dirs:
        next_num = max(existing_dirs) + 1
    
    # Create the new directory
    new_dir = f"Report_{next_num:02d}"
    full_path = os.path.join(base_dir, new_dir)
    os.makedirs(full_path, exist_ok=True)
    
    return full_path

# Create results directory with incremental numbering
def create_numbered_results_dir():
    """Create a numbered results directory"""
    # First make sure the Results directory exists
    os.makedirs("./Results", exist_ok=True)
    
    # Check for existing results directories
    existing_dirs = glob.glob('./Results/res_*/')
    
    # Find the highest numbered directory
    max_num = 0
    for dir_path in existing_dirs:
        try:
            dir_num = int(dir_path.split('_')[1].strip('/\\'))
            max_num = max(max_num, dir_num)
        except (IndexError, ValueError):
            continue
    
    # Create new directory with incremented number
    new_dir_num = max_num + 1
    results_dir = f"./Results/res_{new_dir_num:02d}"
    os.makedirs(results_dir, exist_ok=True)
    
    # Log immediate feedback
    print(f"Created results directory: {results_dir}")
    
    return results_dir

# Set up basic logging
def setup_base_logging(log_dir):
    """Set up basic logging to file and console"""
    if not os.path.exists(log_dir):
        os.makedirs(log_dir)
        
    # Create formatter
    formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
    
    # Create file handler
    log_file = os.path.join(log_dir, 'router_analysis.log')
    file_handler = logging.FileHandler(log_file)
    file_handler.setLevel(logging.INFO)
    file_handler.setFormatter(formatter)
    
    # Create console handler
    console_handler = logging.StreamHandler()
    console_handler.setLevel(logging.INFO)
    console_handler.setFormatter(formatter)
    
    # Create logger
    logger = logging.getLogger('query_router')
    logger.setLevel(logging.INFO)
    
    # Remove existing handlers to avoid duplicates
    if logger.hasHandlers():
        logger.handlers.clear()
        
    logger.addHandler(file_handler)
    logger.addHandler(console_handler)
    
    # Test log message
    logger.info(f"Logger initialized. Writing to {log_file}")
    
    return logger

# Create confusion matrix and save it as an image
def create_confusion_matrix(y_true: List[str], y_pred: List[str], class_names: List[str], file_path: str) -> np.ndarray:
    """
    Create a confusion matrix and save it as an image.
    
    Args:
        y_true: List of true labels
        y_pred: List of predicted labels
        class_names: List of class names
        file_path: Path to save the confusion matrix image
        
    Returns:
        Confusion matrix as numpy array
    """
    cm = confusion_matrix(y_true, y_pred, labels=class_names)
    
    plt.figure(figsize=(10, 8))
    sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', xticklabels=class_names, yticklabels=class_names)
    plt.xlabel('Predicted')
    plt.ylabel('True')
    plt.title('Confusion Matrix')
    plt.tight_layout()
    plt.savefig(file_path)
    plt.close()  # Close the figure to avoid display popup
    
    return cm

# Calculate metrics
def calculate_metrics(y_true: List[str], y_pred: List[str], class_names: List[str]) -> Dict:
    """
    Calculate accuracy, precision, recall, and F1 score.
    
    Args:
        y_true: List of true labels
        y_pred: List of predicted labels
        class_names: List of class names
        
    Returns:
        Dictionary with calculated metrics
    """
    # Calculate accuracy
    accuracy = accuracy_score(y_true, y_pred)
    
    # Calculate precision, recall, and F1 score (weighted average)
    precision, recall, f1, _ = precision_recall_fscore_support(
        y_true, y_pred, average='weighted', zero_division=0)
    
    metrics = {
        'accuracy': accuracy,
        'precision': precision,
        'recall': recall,
        'f1_score': f1
    }
    
    return metrics

# Export results to Excel
def export_results_to_excel(category_errors: List[Dict], subcategory_errors: List[Dict], 
                         correct_items: List[Dict], model_name: str, results_dir: str) -> str:
    """
    Export all classification results to an Excel file.
    
    Args:
        category_errors: List of entries with incorrect category predictions
        subcategory_errors: List of entries with correct category but incorrect subcategory
        correct_items: List of entries with correct category and subcategory
        model_name: Name of the model
        results_dir: Directory to save results
        
    Returns:
        Path to the created Excel file
    """
    # Setup Excel file
    file_name = f"Results for model {model_name.replace('/', '-')}.xlsx"
    file_path = os.path.join(results_dir, file_name)
    logging.getLogger('query_router').info(f"Creating Excel file: {file_path}")
    
    # Create the DataFrame
    data = []
    
    # Add category errors
    for item in category_errors:
        data.append({
            'CResult': 'IC',
            'Sentence': item['query'],
            '1st Probability Category/Subcategory': f"{item['similarity']*100:.1f}% {item['predicted_category']}/{item['predicted_subcategory']}",
            'Sentence for 1st Category/Subcategory': item['predicted_description'],
            '2nd Probability Category/Subcategory': f"{item['second_best_similarity']*100:.1f}% {item['second_best_category']}/{item['second_best_subcategory']}",
            'Sentence for 2nd Category/Subcategory': item['second_best_description'],
            'Gap (1-2)': f"{item['cosine_distance']*100:.1f}%",
            'Def. (1-2)': f"{item['topic_vectors_similarity']*100:.1f}%",
            'True Category': item['true_category'],
            'True Subcategory': item['true_subcategory']
        })
    
    # Add subcategory errors
    for item in subcategory_errors:
        data.append({
            'CResult': 'CCIS',
            'Sentence': item['query'],
            '1st Probability Category/Subcategory': f"{item['similarity']*100:.1f}% {item['predicted_category']}/{item['predicted_subcategory']}",
            'Sentence for 1st Category/Subcategory': item['predicted_description'],
            '2nd Probability Category/Subcategory': f"{item['second_best_similarity']*100:.1f}% {item['second_best_category']}/{item['second_best_subcategory']}",
            'Sentence for 2nd Category/Subcategory': item['second_best_description'],
            'Gap (1-2)': f"{item['cosine_distance']*100:.1f}%",
            'Def. (1-2)': f"{item['topic_vectors_similarity']*100:.1f}%",
            'True Category': item['true_category'],
            'True Subcategory': item['true_subcategory']
        })
    
    # Add correct classifications
    for item in correct_items:
        data.append({
            'CResult': 'CCCS',
            'Sentence': item['query'],
            '1st Probability Category/Subcategory': f"{item['similarity']*100:.1f}% {item['predicted_category']}/{item['predicted_subcategory']}",
            'Sentence for 1st Category/Subcategory': item['predicted_description'],
            '2nd Probability Category/Subcategory': f"{item['second_best_similarity']*100:.1f}% {item['second_best_category']}/{item['second_best_subcategory']}",
            'Sentence for 2nd Category/Subcategory': item['second_best_description'],
            'Gap (1-2)': f"{item['cosine_distance']*100:.1f}%",
            'Def. (1-2)': f"{item['topic_vectors_similarity']*100:.1f}%",
            'True Category': item['true_category'],
            'True Subcategory': item['true_subcategory']
        })
    
    # Create DataFrame
    df = pd.DataFrame(data)
    
    # Sort the data by CResult, then 1st Probability, then 2nd Probability
    # First, extract numeric values from the probability columns for sorting
    df['1st_prob_value'] = df['1st Probability Category/Subcategory'].str.extract(r'([\d.]+)').astype(float)
    df['2nd_prob_value'] = df['2nd Probability Category/Subcategory'].str.extract(r'([\d.]+)').astype(float)
    
    # Sort by CResult, 1st probability (descending), 2nd probability (descending)
    df = df.sort_values(by=['CResult', '1st_prob_value', '2nd_prob_value'], 
                       ascending=[True, False, False])
    
    # Remove temporary columns used for sorting
    df = df.drop(columns=['1st_prob_value', '2nd_prob_value'])
    
    # Group by True Category
    categories = df['True Category'].unique()
    
    # Create an Excel writer with xlsxwriter engine
    with pd.ExcelWriter(file_path, engine='xlsxwriter') as writer:
        # For each category, create a sheet
        for category in categories:
            category_df = df[df['True Category'] == category]
            sheet_name = category[:31]  # Excel sheet names have a max length of 31 characters
            category_df.to_excel(writer, sheet_name=sheet_name, index=False)
            
            # Access the workbook and the worksheet objects
            workbook = writer.book
            worksheet = writer.sheets[sheet_name]
            
            # Apply filter to all columns
            worksheet.autofilter(0, 0, len(category_df), len(df.columns) - 1)
    
    logging.getLogger('query_router').info(f"Exported results to Excel: {file_path}")
    return file_path

# Helper to truncate text
def truncate_text(text: str, max_length: int) -> str:
    """
    Truncate text to a specified maximum length.
    
    Args:
        text: Text to truncate
        max_length: Maximum length of the truncated text
        
    Returns:
        Truncated text
    """
    if len(text) <= max_length:
        return text
    else:
        return text[:max_length-3] + "..."

# Find close categories
def find_close_categories(topics_df: pd.DataFrame, topic_embeddings: np.ndarray, threshold: float = 0.90) -> Tuple[pd.DataFrame, str]:
    """
    Find categories and subcategories that have very close vector representations.
    
    Args:
        topics_df: DataFrame with topic information
        topic_embeddings: Embeddings of the topic descriptions
        threshold: Similarity threshold to consider topics as "close" (default: 0.90)
        
    Returns:
        Tuple containing:
        - DataFrame with pairs of close categories
        - String message about the threshold used for the report
    """
    close_pairs = []
    threshold_message = f"Using a similarity threshold of {threshold*100:.0f}% to identify closely related categories."
    logging.getLogger('query_router').info(threshold_message)
    
    # Calculate similarity between all pairs of topic embeddings
    for i in range(len(topic_embeddings)):
        for j in range(i + 1, len(topic_embeddings)):  # Only compare each pair once
            # Calculate cosine similarity
            similarity = cosine_similarity([topic_embeddings[i]], [topic_embeddings[j]])[0][0]
            
            # If similarity is above threshold, add to results
            if similarity > threshold:
                topic_i = topics_df.iloc[i]
                topic_j = topics_df.iloc[j]
                
                close_pairs.append({
                    'Category1': topic_i['category'],
                    'Subcategory1': topic_i['title'],
                    'Description1': topic_i['description'],
                    'Category2': topic_j['category'],
                    'Subcategory2': topic_j['title'],
                    'Description2': topic_j['description'],
                    'Similarity': similarity,
                    # New formatted output as requested
                    'Formatted_Output': f"{similarity*100:.1f}%   @@   {topic_i['category']}/{topic_i['title']}   ##   {topic_j['category']}/{topic_j['title']}\n{topic_i['description']}   %%   {topic_j['description']}"
                })
    
    # Convert to DataFrame and sort by similarity (descending)
    if close_pairs:
        close_pairs_df = pd.DataFrame(close_pairs)
        close_pairs_df = close_pairs_df.sort_values(by='Similarity', ascending=False)
        return close_pairs_df, threshold_message
    else:
        return pd.DataFrame(columns=['Category1', 'Subcategory1', 'Description1', 
                                    'Category2', 'Subcategory2', 'Description2', 
                                    'Similarity', 'Formatted_Output']), threshold_message

# Functions for error tracking
def get_category_errors(results: List[Dict], true_categories: List[str], true_subcategories: List[str], 
                         topics_df: pd.DataFrame) -> List[Dict]:
    """
    Get entries where the predicted category is wrong.
    """
    category_errors = []
    
    for i, result in enumerate(results):
        if result['predicted_category'] != true_categories[i]:
            # Get the description of the true category/subcategory
            true_desc_row = topics_df[(topics_df['category'] == true_categories[i]) & 
                                     (topics_df['title'] == true_subcategories[i])]
            
            true_description = ""
            if not true_desc_row.empty:
                true_description = true_desc_row.iloc[0]['description']
            
            # Get the description of the second best category/subcategory
            second_best_desc_row = topics_df[(topics_df['category'] == result['second_best_category']) & 
                                           (topics_df['title'] == result['second_best_subcategory'])]
            
            second_best_description = ""
            if not second_best_desc_row.empty:
                second_best_description = second_best_desc_row.iloc[0]['description']
            
            category_errors.append({
                'query': result['query'],
                'true_category': true_categories[i],
                'true_subcategory': true_subcategories[i],
                'true_description': true_description,
                'predicted_category': result['predicted_category'],
                'predicted_subcategory': result['predicted_subcategory'],
                'predicted_description': result['predicted_description'],
                'similarity': result['similarity'],
                'second_best_category': result['second_best_category'],
                'second_best_subcategory': result['second_best_subcategory'],
                'second_best_description': second_best_description,
                'second_best_similarity': result['second_best_similarity'],
                'cosine_distance': result['cosine_distance'],
                'topic_vectors_similarity': result['topic_vectors_similarity']
            })
    
    return category_errors

def get_subcategory_errors(results: List[Dict], true_categories: List[str], true_subcategories: List[str], 
                           topics_df: pd.DataFrame) -> List[Dict]:
    """
    Get entries where the predicted category is correct but subcategory is wrong.
    """
    subcategory_errors = []
    
    for i, result in enumerate(results):
        if (result['predicted_category'] == true_categories[i] and 
            result['predicted_subcategory'] != true_subcategories[i]):
            
            # Get the description of the true category/subcategory
            true_desc_row = topics_df[(topics_df['category'] == true_categories[i]) & 
                                     (topics_df['title'] == true_subcategories[i])]
            
            true_description = ""
            if not true_desc_row.empty:
                true_description = true_desc_row.iloc[0]['description']
            
            # Get the description of the second best category/subcategory
            second_best_desc_row = topics_df[(topics_df['category'] == result['second_best_category']) & 
                                           (topics_df['title'] == result['second_best_subcategory'])]
            
            second_best_description = ""
            if not second_best_desc_row.empty:
                second_best_description = second_best_desc_row.iloc[0]['description']
            
            subcategory_errors.append({
                'query': result['query'],
                'true_category': true_categories[i],
                'true_subcategory': true_subcategories[i],
                'true_description': true_description,
                'predicted_category': result['predicted_category'],
                'predicted_subcategory': result['predicted_subcategory'],
                'predicted_description': result['predicted_description'],
                'similarity': result['similarity'],
                'second_best_category': result['second_best_category'],
                'second_best_subcategory': result['second_best_subcategory'],
                'second_best_description': second_best_description,
                'second_best_similarity': result['second_best_similarity'],
                'cosine_distance': result['cosine_distance'],
                'topic_vectors_similarity': result['topic_vectors_similarity']
            })
    
    return subcategory_errors

def get_correct_classifications(results: List[Dict], true_categories: List[str], true_subcategories: List[str], 
                             topics_df: pd.DataFrame) -> List[Dict]:
    """
    Get entries where both category and subcategory predictions are correct.
    """
    correct_items = []
    
    for i, result in enumerate(results):
        if (result['predicted_category'] == true_categories[i] and 
            result['predicted_subcategory'] == true_subcategories[i]):
            
            # Get the description of the true category/subcategory
            true_desc_row = topics_df[(topics_df['category'] == true_categories[i]) & 
                                     (topics_df['title'] == true_subcategories[i])]
            
            true_description = ""
            if not true_desc_row.empty:
                true_description = true_desc_row.iloc[0]['description']
            
            # Get the description of the second best category/subcategory
            second_best_desc_row = topics_df[(topics_df['category'] == result['second_best_category']) & 
                                           (topics_df['title'] == result['second_best_subcategory'])]
            
            second_best_description = ""
            if not second_best_desc_row.empty:
                second_best_description = second_best_desc_row.iloc[0]['description']
            
            correct_items.append({
                'query': result['query'],
                'true_category': true_categories[i],
                'true_subcategory': true_subcategories[i],
                'true_description': true_description,
                'predicted_category': result['predicted_category'],
                'predicted_subcategory': result['predicted_subcategory'],
                'predicted_description': result['predicted_description'],
                'similarity': result['similarity'],
                'second_best_category': result['second_best_category'],
                'second_best_subcategory': result['second_best_subcategory'],
                'second_best_description': second_best_description,
                'second_best_similarity': result['second_best_similarity'],
                'cosine_distance': result['cosine_distance'],
                'topic_vectors_similarity': result['topic_vectors_similarity']
            })
    
    return correct_items

class OutputLogger:
    """
    A class to redirect stdout to both terminal and a text file.
    """
    def __init__(self, report_dir=None):
        # Create or use the specified report directory
        self.report_dir = report_dir if report_dir else get_next_report_dir()
        
        # Set up the file path
        self.log_file_path = os.path.join(self.report_dir, "stdout.txt")
        
        # Store the original stdout
        self.terminal = sys.stdout
        
        # Open the log file
        self.log = open(self.log_file_path, "w", encoding="utf-8")
        
        # Report data storage
        self.report_data = []
        
    def write(self, message):
        self.terminal.write(message)
        self.log.write(message)
        self.log.flush()
        
    def flush(self):
        self.terminal.flush()
        self.log.flush()
        
    def close(self):
        if self.log:
            self.log.close()
            
    def __del__(self):
        self.close()

    def add_to_report(self, content_type, content, **kwargs):
        """
        Add content to the report data structure
        
        Parameters:
        - content_type: "heading1", "heading2", "text", "image", "table", etc.
        - content: The actual content (text, image path, etc.)
        - kwargs: Additional arguments specific to the content type
        """
        self.report_data.append({
            "type": content_type,
            "content": content,
            "options": kwargs
        })
        
        # If it's regular text, also print to stdout and log
        if content_type == "text":
            self.write(content + "\n")

    def get_report_dir(self):
        """
        Return the current report directory
        """
        return self.report_dir


class ExcelReportGenerator:
    """
    Class to generate Excel reports
    """
    def __init__(self, report_dir):
        self.report_dir = report_dir
        self.workbook = None
        self.worksheet = None
        self.current_row = 1  # Start from row 1 (after header)
        
    def create_excel(self, filename="report.xlsx", sheet_name="Sheet1"):
        """
        Create a new Excel workbook and worksheet
        """
        self.workbook = openpyxl.Workbook()
        self.worksheet = self.workbook.active
        self.worksheet.title = sheet_name
        self.filepath = os.path.join(self.report_dir, filename)
        self.current_row = 1
        return self.worksheet
    
    def add_header_row(self, headers):
        """
        Add a header row to the Excel file
        """
        if not self.worksheet:
            raise ValueError("Worksheet not created. Call create_excel() first.")
        
        for col_idx, header in enumerate(headers, 1):
            self.worksheet.cell(row=1, column=col_idx, value=header)
        
        # Make headers bold
        for cell in self.worksheet[1]:
            cell.font = openpyxl.styles.Font(bold=True)
        
        # Set autofilter
        self.worksheet.auto_filter.ref = f"A1:{openpyxl.utils.get_column_letter(len(headers))}{1}"
        
        self.current_row = 2  # Next row will be 2
        return self.worksheet
    
    def add_row(self, row_data):
        """
        Add a row of data to the Excel file
        """
        if not self.worksheet:
            raise ValueError("Worksheet not created. Call create_excel() first.")
        
        for col_idx, value in enumerate(row_data, 1):
            self.worksheet.cell(row=self.current_row, column=col_idx, value=value)
        
        self.current_row += 1
        return self.worksheet
    
    def apply_filter(self, column_index, values_to_keep):
        """
        Filter the Excel file to only show rows where the specified column contains one of the values in values_to_keep
        """
        if not self.worksheet:
            raise ValueError("Worksheet not created. Call create_excel() first.")
        
        # Get column letter
        col_letter = openpyxl.utils.get_column_letter(column_index)
        
        # Define filter criteria
        filter_column = FilterColumn(colId=column_index-1)  # 0-based index
        filter_column.filters = Filters()
        filter_column.filters.filter = values_to_keep
        
        # Apply filter
        self.worksheet.auto_filter.filterColumn = [filter_column]
        
        return self.worksheet
    
    def sort_by_column(self, column_index, descending=False):
        """
        Sort the Excel file by the specified column
        Note: This will sort the data in memory. The actual sorting in Excel
        will need to be applied by the user after opening the file.
        """
        if not self.worksheet:
            raise ValueError("Worksheet not created. Call create_excel() first.")
        
        # Get all data excluding header
        data = []
        for row in self.worksheet.iter_rows(min_row=2, values_only=True):
            data.append(row)
        
        # Sort data
        data.sort(key=lambda x: x[column_index-1], reverse=descending)
        
        # Clear existing data
        for row in range(2, self.current_row):
            for col in range(1, self.worksheet.max_column + 1):
                self.worksheet.cell(row=row, column=col).value = None
        
        # Write sorted data back
        self.current_row = 2
        for row_data in data:
            for col_idx, value in enumerate(row_data, 1):
                self.worksheet.cell(row=self.current_row, column=col_idx, value=value)
            self.current_row += 1
        
        return self.worksheet
    
    def save_and_close(self):
        """
        Save the Excel file and close the workbook
        """
        if not self.workbook:
            raise ValueError("Workbook not created. Call create_excel() first.")
        
        self.workbook.save(self.filepath)
        print(f"Excel report saved to {self.filepath}")
        return self.filepath


class ReportGenerator:
    """
    Class to generate reports from the logged data and additional content.
    """
    def __init__(self, logger=None, report_directory=None):
        if logger:
            self.logger = logger
            self.report_directory = logger.get_report_dir()
        elif report_directory:
            self.report_directory = report_directory
            self.logger = None
        else:
            self.report_directory = get_next_report_dir()
            self.logger = None
        
        # Track if we're in two-column mode
        self.two_column_mode = False
        self.current_column = 0  # 0: Not in column mode, 1: Column 1, 2: Column 2
        
        # For storing document and sections
        self.current_doc = None
        self.main_section = None
        self.column_section = None
        
        # Create a logger for this report generator
        self.log = logging.getLogger('report_generator')
        
    def _create_front_page(self, doc, title, subtitle=None, author=None, organization=None):
        """Create a professional front page for the report"""
        section = doc.sections[0]
        
        # Add title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(24)
        
        # Add subtitle if provided
        if subtitle:
            subtitle_para = doc.add_paragraph()
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle_para.add_run(subtitle)
            subtitle_run.italic = True
            subtitle_run.font.size = Pt(16)
        
        # Add date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(datetime.datetime.now().strftime("%B %d, %Y"))
        date_run.font.size = Pt(12)
        
        # Add author if provided
        if author:
            author_para = doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_para.add_run(f"\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\nPrepared by: {author}")
        
        # Add organization if provided
        if organization:
            org_para = doc.add_paragraph()
            org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            org_para.add_run(organization)
            
        # Add page break after cover page
        doc.add_page_break()
    
    def _insert_toc(self, doc):
        """Insert table of contents"""
        doc.add_paragraph("Table of Contents").style = 'Heading 1'
        
        # Add TOC field
        doc.add_paragraph().add_run("TOC WILL BE GENERATED AUTOMATICALLY").bold = True
        
        # In a real Word document, you would use this field code:
        # {TOC \o "1-2" \h \z \u}
        # But we use a placeholder as python-docx doesn't directly support field codes
        
        doc.add_page_break()
    
    def _insert_external_content(self, doc, file_path):
        """Insert content from an external document"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                doc.add_paragraph(content)
        except Exception as e:
            doc.add_paragraph(f"Error including external content: {str(e)}")
            
    def _convert_svg_to_png(self, svg_path):
        """Convert SVG to PNG for Word compatibility"""
        if not os.path.exists(svg_path):
            print(f"SVG file not found: {svg_path}")
            return None
        
        # Get the base filename without extension
        base_name = os.path.splitext(os.path.basename(svg_path))[0]
        png_path = os.path.join(self.report_directory, f"{base_name}.png")
        
        try:
            if HAS_CAIROSVG:
                # Use cairosvg to convert SVG to PNG
                with open(svg_path, 'rb') as svg_file:
                    svg_content = svg_file.read()
                    svg2png(bytestring=svg_content, write_to=png_path, output_width=800)
                return png_path
            elif HAS_PIL:
                # Fallback warning
                print("Warning: CairoSVG not found. SVG conversion may be suboptimal.")
                # Try using other methods (this is a placeholder - real implementation would need a different approach)
                # This won't actually work but shows where to integrate alternate conversion
                return None
            else:
                print("Warning: SVG conversion libraries not found. Skipping SVG image.")
                return None
        except Exception as e:
            print(f"Error converting SVG to PNG: {str(e)}")
            return None
    
    def create_document(self, output_filename="report.docx", paper_format="A4", orientation="landscape"):
        """
        Create a new Word document with the specified format and orientation
        
        Parameters:
        - output_filename: Name of the output file
        - paper_format: "A4", "Letter", or "Legal"
        - orientation: "landscape" or "portrait"
        
        Returns:
        - The Document object
        """
        # Create document
        doc = Document()
        self.current_doc = doc
        
        # Get the default section
        section = doc.sections[0]
        self.main_section = section
        
        # Set orientation
        if orientation.lower() == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
        else:
            section.orientation = WD_ORIENT.PORTRAIT
        
        # Set paper format
        if paper_format.upper() == "A4":
            section.page_width = Mm(297)
            section.page_height = Mm(210)
        elif paper_format.upper() == "LETTER":
            section.page_width = Inches(11)
            section.page_height = Inches(8.5)
        elif paper_format.upper() == "LEGAL":
            section.page_width = Inches(14)
            section.page_height = Inches(8.5)
        
        # Adjust page size if in portrait mode
        if orientation.lower() == "portrait":
            section.page_width, section.page_height = section.page_height, section.page_width
        
        # Store the file path
        self.output_path = os.path.join(self.report_directory, output_filename)
        
        return doc
    
    def start_two_column_mode(self):
        """
        Start a two-column layout in the Word document
        """
        if not self.current_doc:
            raise ValueError("Document not created. Call create_document() first.")
        
        # Create a new section for the two columns
        new_section = self.current_doc.add_section(WD_SECTION.NEW_PAGE)
        self.column_section = new_section
        
        # Set orientation and page size to match the main section
        new_section.orientation = self.main_section.orientation
        new_section.page_width = self.main_section.page_width
        new_section.page_height = self.main_section.page_height
        
        # Set two columns
        new_section.left_margin = self.main_section.left_margin
        new_section.right_margin = self.main_section.right_margin
        new_section.top_margin = self.main_section.top_margin
        new_section.bottom_margin = self.main_section.bottom_margin
        new_section.column_count = 2
        
        self.two_column_mode = True
        self.current_column = 1
        
        return self.current_doc
    
    def save_document(self):
        """
        Save the document
        """
        if not self.current_doc:
            raise ValueError("Document not created. Call create_document() first.")
        
        self.current_doc.save(self.output_path)
        print(f"Document saved to {self.output_path}")
        return self.output_path
    
    def generate_report(self, output_filename="report.docx", title="ML/NLP Research Report", 
                        subtitle="Results and Analysis", author=None, organization=None, intro_file=None,
                        paper_format="A4", orientation="landscape"):
        """
        Generate a professional report based on collected data
        """
        # Create document with the specified format
        doc = self.create_document(output_filename, paper_format, orientation)
        
        # Front page
        self._create_front_page(doc, title, subtitle, author, organization)
        
        # Table of contents
        self._insert_toc(doc)
        
        # Introduction section
        doc.add_heading("Introduction", level=1)
        if intro_file and os.path.exists(intro_file):
            self._insert_external_content(doc, intro_file)
        else:
            doc.add_paragraph("Introduction content not found or not specified.")
        
        # Process all the content collected in the logger
        if self.logger:
            for item in self.logger.report_data:
                content_type = item["type"]
                content = item["content"]
                options = item["options"]
                
                if content_type == "heading1":
                    doc.add_heading(content, level=1)
                
                elif content_type == "heading2":
                    doc.add_heading(content, level=2)
                
                elif content_type == "text":
                    paragraph = doc.add_paragraph(content)
                    
                    # Apply formatting options if provided
                    if options.get("bold", False):
                        for run in paragraph.runs:
                            run.bold = True
                    if options.get("italic", False):
                        for run in paragraph.runs:
                            run.italic = True
                
                elif content_type == "image":
                    # Check if the file exists
                    if os.path.exists(content):
                        paragraph = doc.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        width = options.get("width", 6)  # Default width in inches
                        
                        # Check if it's an SVG file
                        if content.lower().endswith('.svg'):
                            # Try to convert SVG to PNG
                            png_path = self._convert_svg_to_png(content)
                            if png_path and os.path.exists(png_path):
                                doc.add_picture(png_path, width=Inches(width))
                            else:
                                # If conversion fails, add a placeholder text
                                doc.add_paragraph(f"[SVG image: {os.path.basename(content)}]")
                                print(f"Warning: Could not convert SVG to PNG: {content}")
                        else:
                            # For non-SVG images
                            try:
                                doc.add_picture(content, width=Inches(width))
                            except Exception as e:
                                doc.add_paragraph(f"[Image could not be added: {os.path.basename(content)}]")
                                print(f"Error adding image: {str(e)}")
                    else:
                        doc.add_paragraph(f"[Image not found: {content}]")
                
                elif content_type == "table":
                    # Add tables
                    table = self.current_doc.add_table(rows=len(content), cols=len(content[0]) if content else 0)
                    table.style = 'Table Grid'
                    for i, row_data in enumerate(content):
                        for j, cell_data in enumerate(row_data):
                            table.cell(i, j).text = str(cell_data)
        
        # Save the document
        return self.save_document()

# Function to set up the logger and reporter
def setup_logging_and_reporting(report_dir=None):
    """
    Set up the logging and reporting system
    
    Returns a tuple containing:
    - logger: The OutputLogger instance
    - reporter: The ReportGenerator instance
    - excel_reporter: The ExcelReportGenerator instance
    - report_dir: The path to the report directory
    - basic_logger: The basic logger instance
    """
    # If no report directory provided, get the next available one
    if not report_dir:
        report_dir = create_numbered_results_dir()
    
    # Set up basic logging
    basic_logger = setup_base_logging(report_dir)
    
    # Set up output logger
    logger = OutputLogger(report_dir)
    # Redirect stdout to our logger
    sys.stdout = logger
    
    # Create the reporters
    reporter = ReportGenerator(logger, report_dir)
    excel_reporter = ExcelReportGenerator(report_dir)
    
    return logger, reporter, excel_reporter, report_dir, basic_logger

# For backwards compatibility
from docx.enum.text import WD_BREAK
    
def add_heading(self, text, level=1):
    """
    Add a heading to the document
    """
    if not self.current_doc:
        raise ValueError("Document not created. Call create_document() first.")
    
    self.current_doc.add_heading(text, level=level)
    return self.current_doc
    
def add_paragraph(self, text, **kwargs):
    """
    Add a paragraph to the document
    
    Parameters:
    - text: The paragraph text
    - kwargs: Additional formatting options (bold, italic, etc.)
    """
    if not self.current_doc:
        raise ValueError("Document not created. Call create_document() first.")
    
    para = self.current_doc.add_paragraph(text)
    
    # Apply formatting
    if kwargs.get("bold", False):
        for run in para.runs:
            run.bold = True
    if kwargs.get("italic", False):
        for run in para.runs:
            run.italic = True
    
    return self.current_doc

def add_image(self, image_path, width=6):
    """
    Add an image to the document
    
    Parameters:
    - image_path: Path to the image file
    - width: Width in inches
    """
    if not self.current_doc:
        raise ValueError("Document not created. Call create_document() first.")
    
    # Check if the file exists
    if not os.path.exists(image_path):
        self.current_doc.add_paragraph(f"[Image not found: {image_path}]")
        return self.current_doc
    
    # Check if it's an SVG file
    if image_path.lower().endswith('.svg'):
        # Try to convert SVG to PNG
        png_path = self._convert_svg_to_png(image_path)
        if png_path and os.path.exists(png_path):
            self.current_doc.add_picture(png_path, width=Inches(width))
        else:
            # If conversion fails, add a placeholder text
            self.current_doc.add_paragraph(f"[SVG image: {os.path.basename(image_path)}]")
    else:
        # For non-SVG images
        try:
            self.current_doc.add_picture(image_path, width=Inches(width))
        except Exception as e:
            self.current_doc.add_paragraph(f"[Image could not be added: {os.path.basename(image_path)}]")
            print(f"Error adding image: {str(e)}")
    
    return self.current_doc

def add_table(self, table_data):
    """
    Add a table to the document
    
    Parameters:
    - table_data: A list of lists representing the table data
    """
    if not self.current_doc:
        raise ValueError("Document not created. Call create_document() first.")
    
    if not table_data or not isinstance(table_data, list) or len(table_data) == 0:
        return self.current_doc
    
    # Create table with appropriate dimensions
    num_rows = len(table_data)
    num_cols = len(table_data[0]) if num_rows > 0 else 0
    
    table = self.current_doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    
    # Fill the table with data
    for i, row_data in enumerate(table_data):
        for j, cell_data in enumerate(row_data):
            table.cell(i, j).text = str(cell_data)
    
    return self.current_doc

def write_in_column_1(self):
    """
    Switch to writing in column 1
    """
    if not self.two_column_mode:
        raise ValueError("Not in two-column mode. Call start_two_column_mode() first.")
    
    self.current_column = 1
    return self.current_doc

def write_in_column_2(self):
    """
    Switch to writing in column 2
    """
    if not self.two_column_mode:
        raise ValueError("Not in two-column mode. Call start_two_column_mode() first.")
    
    self.current_column = 2
    
    # In python-docx we can't directly write to a specific column
    # We use a column break to move to the next column
    if self.current_column == 2:
        self.current_doc.add_paragraph().add_run().add_break(WD_BREAK.COLUMN)
    
    return self.current_doc

def stop_two_column_mode(self):
    """
    End the two-column layout and return to normal
    """
    if not self.two_column_mode:
        return self.current_doc
    
    # Add a new section to return to single column
    new_section = self.current_doc.add_section(WD_SECTION.NEW_PAGE)
    
    # Copy settings from main section
    new_section.orientation = self.main_section.orientation
    new_section.page_width = self.main_section.page_width
    new_section.page_height = self.main_section.page_height
    new_section.left_margin = self.main_section.left_margin
    new_section.right_margin = self.main_section.right_margin
    new_section.top_margin = self.main_section.top_margin
    new_section.bottom_margin = self.main_section.bottom_margin
    new_section.column_count = 1
    
    self.two_column_mode = False
    self.current_column = 0
    self.main_section = new_section
    
    return self.current_doc
